from __future__ import annotations

import re
from pathlib import Path
from urllib.parse import unquote

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image


REPO_ROOT = Path(__file__).resolve().parent.parent
REPORT_MD = REPO_ROOT / "REPORT.md"
OUTPUT_DOCX = REPO_ROOT / "outputs" / "MATH5380 Report.docx"


INLINE_TOKEN_RE = re.compile(
    r"(\[[^\]]+\]\([^)]+\)|\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)"
)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)

    new_run.append(r_pr)
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def resolve_link(target: str, report_dir: Path) -> str | None:
    if target.startswith("#"):
        return None
    if "://" in target:
        return target
    local_target = report_dir / unquote(target)
    if local_target.exists():
        return local_target.resolve().as_uri()
    return target


def add_inline_text(paragraph, text: str, report_dir: Path) -> None:
    pos = 0
    for match in INLINE_TOKEN_RE.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])

        token = match.group(0)
        if token.startswith("[") and "](" in token:
            label, target = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token).groups()
            resolved = resolve_link(target, report_dir)
            if resolved:
                add_hyperlink(paragraph, label, resolved)
            else:
                paragraph.add_run(label)
        elif token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Menlo"
            run.font.size = Pt(10)
        elif token.startswith("*") and token.endswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        else:
            paragraph.add_run(token)

        pos = match.end()

    if pos < len(text):
        paragraph.add_run(text[pos:])


def parse_table(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        stripped = line.strip()
        if not stripped.startswith("|"):
            continue
        cells = [cell.strip() for cell in stripped.strip("|").split("|")]
        rows.append(cells)
    return rows


def add_markdown_table(document: Document, table_lines: list[str], report_dir: Path) -> None:
    rows = parse_table(table_lines)
    if len(rows) < 2:
        return
    data_rows = [rows[0]] + rows[2:]
    table = document.add_table(rows=len(data_rows), cols=len(data_rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for r_idx, row in enumerate(data_rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            add_inline_text(paragraph, cell_text, report_dir)
            if r_idx == 0:
                for run in paragraph.runs:
                    run.bold = True


def add_image(document: Document, image_target: str, alt_text: str, report_dir: Path) -> None:
    image_path = (report_dir / unquote(image_target)).resolve()
    if not image_path.exists():
        paragraph = document.add_paragraph()
        paragraph.add_run(f"[Missing image: {alt_text}]").italic = True
        return

    max_width = 6.1
    with Image.open(image_path) as img:
        width_px, height_px = img.size
    aspect = height_px / width_px if width_px else 0.75
    width_inches = min(max_width, max(4.3, width_px / 200))
    height_inches = width_inches * aspect

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches), height=Inches(height_inches))


def build_docx() -> Path:
    report_dir = REPORT_MD.parent
    lines = REPORT_MD.read_text(encoding="utf-8").splitlines()

    document = Document()
    document.core_properties.title = "GTAA Factor Portfolio — Project 2 Report"
    document.core_properties.author = "MATH 5380 Project Team"
    document.core_properties.subject = "MATH 5380 Project 2"

    section = document.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Aptos"
    normal_style.font.size = Pt(10)
    normal_style.paragraph_format.space_after = Pt(3)
    normal_style.paragraph_format.line_spacing = 1.0

    title_style = document.styles["Title"]
    title_style.font.name = "Aptos Display"
    title_style.font.size = Pt(18)

    for style_name, size in [("Heading 1", 13), ("Heading 2", 11.5), ("Heading 3", 10.5)]:
        style = document.styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(size)
        style.paragraph_format.space_before = Pt(4)
        style.paragraph_format.space_after = Pt(2)

    idx = 0
    while idx < len(lines):
        line = lines[idx].rstrip()
        stripped = line.strip()

        if not stripped:
            idx += 1
            continue

        if stripped == "---":
            idx += 1
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped[level:].strip()
            if idx == 0 and level == 1:
                document.add_paragraph(text, style="Title")
            else:
                document.add_heading(text, level=min(level, 4))
            idx += 1
            continue

        if stripped.startswith("!["):
            match = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
            if match:
                alt_text, image_target = match.groups()
                add_image(document, image_target, alt_text, report_dir)
            idx += 1
            continue

        if stripped.startswith("|") and idx + 1 < len(lines) and set(lines[idx + 1].replace("|", "").replace(":", "").replace("-", "").strip()) == set():
            table_lines = [lines[idx], lines[idx + 1]]
            idx += 2
            while idx < len(lines) and lines[idx].strip().startswith("|"):
                table_lines.append(lines[idx])
                idx += 1
            add_markdown_table(document, table_lines, report_dir)
            continue

        if stripped.startswith("- "):
            while idx < len(lines) and lines[idx].strip().startswith("- "):
                paragraph = document.add_paragraph(style="List Bullet")
                add_inline_text(paragraph, lines[idx].strip()[2:].strip(), report_dir)
                idx += 1
            continue

        if re.match(r"\d+\.\s+", stripped):
            while idx < len(lines) and re.match(r"\d+\.\s+", lines[idx].strip()):
                paragraph = document.add_paragraph(style="List Number")
                item_text = re.sub(r"^\d+\.\s+", "", lines[idx].strip())
                add_inline_text(paragraph, item_text, report_dir)
                idx += 1
            continue

        para_lines = [stripped]
        idx += 1
        while idx < len(lines):
            candidate = lines[idx].strip()
            if not candidate:
                break
            if candidate == "---":
                break
            if candidate.startswith("#"):
                break
            if candidate.startswith("!["):
                break
            if candidate.startswith("|") and idx + 1 < len(lines) and set(lines[idx + 1].replace("|", "").replace(":", "").replace("-", "").strip()) == set():
                break
            if candidate.startswith("- "):
                break
            if re.match(r"\d+\.\s+", candidate):
                break
            para_lines.append(candidate)
            idx += 1

        paragraph = document.add_paragraph()
        add_inline_text(paragraph, " ".join(para_lines), report_dir)

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


if __name__ == "__main__":
    output = build_docx()
    print(output)
